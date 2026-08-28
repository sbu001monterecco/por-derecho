#!/usr/bin/env python3
"""Build the bilingual public-safe Concurso 36/2012 continuity assessment."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "assets/data/concurso36-decision-continuity-2014-2026-v1.json"
OUTPUT = ROOT / "assets/data/concurso36-decision-continuity-assessment-2014-2026.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x5E, 0x6B, 0x78)
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    headings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in headings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run("POR DERECHO · CONCURSO 36/2012^ · CONTINUITY CONTROL")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer.paragraphs[0]
        footer.paragraph_format.space_before = Pt(0)
        add_page_number(footer)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("INTERIM ASSESSMENT / EVALUACIÓN INTERMEDIA")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Concurso 36/2012^\nDecision and AC-report continuity, 2014–2026")
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Continuidad de resoluciones e informes de la Administración Concursal, 2014–2026")
    set_run_font(run, size=13, color=MUTED, italic=True)

    for label, value in (
        ("Control", "PD-C36-DECISION-CONTINUITY-20260828-01"),
        ("Cut-off / Corte", "28 August / agosto 2026"),
        ("Status / Estado", "PARTIAL — certified denominator and procedural families remain open"),
        ("Publication / Publicación", "Local candidate only; not pushed, merged, deployed, emailed or filed"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)


def add_callout(doc, title, text, fill=CALLOUT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    r.add_break()
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, label in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT)
        p = header.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for values in rows:
        new_row = table.add_row()
        prevent_row_split(new_row)
        cells = new_row.cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=9.2, color=INK)
        set_table_geometry(table, widths)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)


def build():
    data = json.loads(DATA.read_text(encoding="utf-8"))
    result = data["result"]
    reports = data["ac_report_continuity"]

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_block(doc)

    doc.add_heading("1. Result / Resultado", level=1)
    add_callout(
        doc,
        "Controlled conclusion / Conclusión controlada",
        "The public-safe audit now reconciles every canonical judicial act in the 2017–2026 complete-record catalogue. It does not establish that the whole court file or every AC report has been obtained. / La auditoría pública segura concilia ya todos los actos judiciales canónicos del catálogo 2017–2026. No acredita que se haya obtenido todo el expediente ni todos los informes de la Administración Concursal.",
    )
    add_table(
        doc,
        ["Class / Clase", "Count / Total", "Control meaning / Significado"],
        [
            ("Core primary decisions / Resoluciones centrales", result["core_primary_decisions_controlled"], "Controlled copy, 2018–2026"),
            ("Earlier in-case anchors / Anclas internas", result["earlier_in_case_anchor_decisions_controlled"], "Controlled copy, 2017"),
            ("Contextual decisions / Resoluciones contextuales", result["connected_or_contextual_primary_decisions_controlled"], "Separate civil/derivative lanes"),
            ("Court-office acts / Actos de oficina judicial", result["controlled_court_office_acts"], "Copy controlled; family incomplete"),
            ("Open or partial families / Familias abiertas", result["unresolved_or_partial_family_rows"], "Primary act or family still incomplete"),
            ("Total audit rows / Filas totales", result["audited_rows"], "Same-date acts remain separate"),
        ],
        [4200, 1200, 3960],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Source note / Nota de fuente: ")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    r = p.add_run("Repository records and connected Google Drive discovery were reviewed. Canonical inventories also preserve prior email-derived source control. No private provider ID, URL, address or filename is reproduced here.")
    set_run_font(r, size=9.5, color=MUTED)

    doc.add_page_break()
    doc.add_heading("2. Material corrections / Correcciones materiales", level=1)
    corrections = [
        ("2017 anchors", "Added C36-E002 and preserved the 22 March, 5 May and 19 December decisions as four earlier in-case anchors."),
        ("8/15 February 2018", "Removed the unsupported second-act row. One C36-E007 act remains under 15 February repository control; body/recital and signature date layers require direct reinspection."),
        ("Controlled omissions", "Restored C36-E009, E030, E035, E039, E040, E042 and R26 from canonical sources."),
        ("Same-date acts", "Separated eight pairs: 16 Apr 2018; 24 Oct 2019; 12 May 2020; 24 Feb 2021; 6 May, 15 Oct 2021; 26 Jan 2022; 12 Sep 2025."),
        ("18 May 2021", "Reports the operative approval wording without treating the copy itself as proof of service, review or procedural finality."),
        ("2022 judicial leads", "Gives separate open rows to E056, E060, E062 and E063 and keeps complaints 375/379 separate."),
        ("Constitutional Court", "Treats the 2811/2023 inadmission date as reported and unconfirmed until the official decision is obtained."),
        ("Separate proceedings", "Keeps Judgment 89/2014 and Judgment 4/2026 contextual and outside the in-docket core count."),
    ]
    add_table(doc, ["Issue / Tema", "Correction / Corrección"], corrections, [2200, 7160])

    doc.add_heading("3. Decision-continuity assessment / Evaluación de resoluciones", level=1)
    add_bullet(doc, "All 39 canonical judicial acts dated 2017–2026 in the complete-record catalogue are represented exactly once in the 51-row projection.")
    add_bullet(doc, "The 28 core decisions are controlled copies in the 2018–2026 Insolvency spine; the four 2017 acts are earlier anchors, and the two contextual decisions remain outside that count.")
    add_bullet(doc, "The 25 October 2021 Diligencia is verified/partial; the signed 27 October 2021 Diligencia is controlled. Their underlying filing, testimony and report families remain incomplete.")
    add_bullet(doc, "The 2014–2016 discovery rows and the year-only 2022 complaints row are declared aggregation exceptions until individual originals can be authenticated.")
    add_bullet(doc, "‘Not located’ means absent from the controlled corpus; it never means nonexistent without a certified denominator.")

    doc.add_page_break()
    doc.add_heading("4. AC-report continuity / Continuidad de informes AC", level=1)
    add_callout(
        doc,
        "Partial series / Serie parcial",
        f"Five report texts are controlled, two court receipts point to report texts not located, and one later/final-report family remains open. / Se controlan cinco textos, dos recepciones judiciales remiten a informes no localizados y queda abierta una familia de informes posteriores/finales.",
        fill="FFF8E8",
    )
    report_state_labels = {
        "REPORT_TEXT_CONTROLLED__ANNEX_AND_BANK_RECONCILIATION_OPEN": "Text controlled; annex/bank reconciliation open / Texto controlado; anexos/banco abiertos",
        "REPORT_TEXT_CONTROLLED__SEPARATE_CLASSIFICATION_LANE": "Text controlled; separate classification lane / Texto controlado; vía de calificación separada",
        "COURT_RECEIPT_CONTROLLED__UNDERLYING_REPORT_NOT_LOCATED": "Court receipt controlled; report not located / Recepción controlada; informe no localizado",
        "REPORT_TEXT_CONTROLLED__POST_DEED_ACCOUNTS_OPEN": "Text controlled; post-deed accounts open / Texto controlado; cuentas posteriores abiertas",
        "REPORT_TEXT_CONTROLLED__INDEPENDENT_ACCOUNTING_PROOF_OPEN": "Text controlled; independent accounting proof open / Texto controlado; prueba contable independiente abierta",
        "SERIES_AND_FINAL_REPORTS_NOT_LOCATED__OFFICIAL_DENOMINATOR_REQUIRED": "Later/final series open; certified denominator required / Serie posterior/final abierta; falta denominador certificado",
    }
    report_rows = []
    for item in reports["records"]:
        report_rows.append((f'{item["date"] or "Undated / Sin fecha"}\n{item["id"]}', f'{item["title_en"]} / {item["title_es"]}', report_state_labels[item["control_state"]], item["control_state"]))
    doc.add_heading("Controlled report texts / Textos controlados", level=2)
    controlled_reports = [values[:3] for values in report_rows if values[3].startswith("REPORT_TEXT_CONTROLLED")]
    add_table(doc, ["Date + ID / Fecha + ID", "Report / Informe", "Control state / Estado"], controlled_reports, [2000, 4300, 3060])
    doc.add_heading("Receipt-only and open records / Recepciones y registros abiertos", level=2)
    open_reports = [values[:3] for values in report_rows if not values[3].startswith("REPORT_TEXT_CONTROLLED")]
    add_table(doc, ["Date + ID / Fecha + ID", "Report / Informe", "Control state / Estado"], open_reports, [2000, 4300, 3060])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("Closure test / Prueba de cierre: ")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    r = p.add_run("certified index of every report and annex from liquidation opening through conclusion; filing/transfer records; challenges and rulings; estate ledger; bank statements; final accounts and conclusion act.")
    set_run_font(r, size=9.5, color=MUTED)

    doc.add_heading("5. Priority gaps / Vacíos prioritarios", level=1)
    gaps = [
        "Certified chronological court/LAJ denominator from 1 January 2014 through conclusion.",
        "Signed originals and complete families for C36-E004, E031, E034, E045, E056, E060, E062 and E063.",
        "Assignment of references 204/14, 677/14, 2418/14 and 5599/2015 to the correct court, class, object and parent proceeding.",
        "Direct reinspection of the C36-E007 date layers without inferring a separate 8 February act.",
        "Underlying filings/testimonies and report 7425/2021 for the 25/27 October 2021 court-office acts.",
        "Official complaint/nullity/Constitutional Court decisions and service/finality for the 2022–2023 review chain.",
        "Merits, service and finality in RPL 2523/2025, combined RPL 3304/2025 and 3319/2025, and RPL 421/2026.",
        "Report 5367/2022, all later reports, final report, final accounts, objections, approval/refusal and conclusion order.",
    ]
    for gap in gaps:
        add_bullet(doc, gap)

    doc.add_heading("6. Repository and website update / Actualización", level=1)
    add_table(
        doc,
        ["Surface / Superficie", "Update / Cambio"],
        [
            ("Canonical JSON", "51-row model, bilingual metadata, typed dates, same-date sequence, public-access state, AC-report supplement and reusable prompt."),
            ("XLSX matrix", "Reproducible workbook with summary, 51-row family matrix and bilingual status legend."),
            ("DOCX/PDF", "This public-safe bilingual interim assessment, rendered and visually checked."),
            ("ES/EN routes", "Dynamic 2014–2026 chronology, public-access labels, report downloads and safe transcript navigation."),
            ("Controls", "Dedicated validation workflow, canonical cross-checks, route-registry discovery and additive maintenance/correction/missing-evidence records."),
        ],
        [2350, 7010],
    )

    doc.add_heading("7. Prepared production request / Solicitud preparada", level=1)
    p = doc.add_paragraph()
    r = p.add_run(data["gap_request_es"])
    set_run_font(r, size=10.5, color=INK)
    p = doc.add_paragraph()
    r = p.add_run(data["gap_request_en"])
    set_run_font(r, size=10.5, color=MUTED, italic=True)
    add_callout(doc, "Authority boundary / Límite de autorización", "This is a prepared specification only. No email, filing, portal submission or authority contact was sent or authorised. / Es sólo una especificación preparada. No se envió ni autorizó correo, escrito, presentación por portal o contacto con autoridad.")

    doc.add_heading("8. Reusable prompt / Prompt reutilizable", level=1)
    p = doc.add_paragraph()
    r = p.add_run(data["audit_prompt"]["es"])
    set_run_font(r, size=10.2, color=INK)
    p = doc.add_paragraph()
    r = p.add_run(data["audit_prompt"]["en"])
    set_run_font(r, size=10.2, color=MUTED, italic=True)

    doc.add_heading("9. Evidential and publication boundary / Límite", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(
        "A decision copy closes only its copy node, not filings, service, review, finality or implementation/accounting. "
        "An AC report proves what was reported, not each underlying financial, bank, title or Registry fact. "
        "No unlawfulness, nonexistence, procedural finality or merits outcome is inferred from an unlocated record. "
        "This candidate is local and unpublished; no push, pull request, merge, deployment, live verification, email or filing is claimed."
    )
    set_run_font(r, size=9.5, color=INK)

    doc.core_properties.title = "Concurso 36/2012 decision and AC-report continuity assessment, 2014–2026"
    doc.core_properties.subject = "Public-safe interim assessment"
    doc.core_properties.keywords = "Concurso 36/2012, continuity audit, court decisions, AC reports"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
